"""control room for the transliterator

from here, open files, get line, feed transliterator each line
 can deal with text files here for now, but docworker
will be the place to pull out lines from .docx

"""

from functools import partial
from io import BytesIO
import mimetypes
from pathlib import Path

from werkzeug.datastructures import FileStorage
from docx import Document as load_docx
from docx.document import Document
from docx.text import paragraph

from hulqcorpustools.transliterator import replaceengine as repl
from hulqcorpustools.resources.constants import TextFormat
from hulqcorpustools.resources.graphemes import Graphemes
from hulqcorpustools.utils.files import FileHandler
from hulqcorpustools.utils.textcounter import TextCounter


class TransliterandFile():
    """_summary_
    """
    def __init__(
            self, 
            _file: FileStorage, 
            source_format: TextFormat, 
            target_format: TextFormat,
            search_method=None,
            out=Path | None,
            **kwargs):

        self.name = Path(_file.filename)

        self._file = _file
        self.source_format = source_format
        self.target_format = target_format
        self.search_method = search_method
        self.out = out
        self.file_type = Path(self.name).suffix

    @property
    def out_name(self):
        return Path((
            f"{self.name.stem} - "
            f"{self.source_format} to {self.target_format}"
            f"{self.name.suffix}"
            ))

    @property
    def out_path(self):
        if self.out is not None:
            out_path = Path(self.out).joinpath(self.out_name)
        else:
            out_path = self.out_name
        
        return out_path

    
    def __fspath__(self):
        return str(self.name)


class DocxTransliterator():

    def __init__(
            self,
            text_counter: TextCounter,
            graphemes: Graphemes):

        self.text_counter = text_counter
        self.graphemes = graphemes

    def _transliterate_paragraph_wordlist(
        self,
        par: paragraph,
        source_format: TextFormat,
        target_format: TextFormat,
        **kwargs):
        '''Transliterates an entire docx file by comparing to as wordlist

        Keyword arguments:
            update_wordlist -- include this if you want to update the wordlist
                            with words found in the transliterated lines
        '''
        new_par_text_parts = []
        # split at tab in case document separates English parts that way
        par_text_parts = par.text.split('\t')

        for par_text_part in par_text_parts:
            if self.text_counter.determine_text_format(par_text_part) != 'english':
                par_text_part = repl.transliterate_string(
                    par_text_part,
                    source_format,
                    target_format,
                    self.graphemes
                )
            new_par_text_parts.append(par_text_part)

            # reassemble tab-separated parts
            par.text = '\t'.join(new_par_text_parts)

    def _transliterate_paragraph_font(
        self,
        par: paragraph,
        target_format: TextFormat
        ):

        if par.style.font.name == 'Straight':
            par.text = repl.transliterate_string(
                par.text,
                TextFormat("straight"),
                target_format,
                self.graphemes)

        else:
            for run in par.runs:
                if run.font.name == 'Straight':
                    run.text = repl.transliterate_string(
                        run.text,
                        TextFormat("straight"),
                        target_format,
                        self.graphemes
                    )

    def transliterate(
        self,
        transliterand: TransliterandFile,
        **kwargs
        ) -> Path:
        """delegate transliteration based on wordlist or font search

        Args:
            transliterand (TransliterandFile): _description_
            source_format (TextFormat): _description_
            target_format (TextFormat): _description_
        """
        _docx = load_docx(transliterand._file)  # type: Document
        if kwargs.get("search_format") == "font":
            _tr_fn = partial(
                self._transliterate_paragraph_font,
                target_format=transliterand.target_format
            )
        else:
            _tr_fn = partial(
                self._transliterate_paragraph_wordlist,
                source_format=transliterand.source_format,
                target_format=transliterand.target_format
            )

        for par in _docx.paragraphs:
            _tr_fn(par)

        for table in _docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        _tr_fn(par)

        _docx.save(transliterand.out_path)

        return transliterand.out_path


class TxtTransliterator():

    def __init__(self, text_counter: TextCounter, graphemes: Graphemes):
        ...

    def transliterate_txt_wordlist(
        txt_transliterand: Path,
        source_format: TextFormat,
        target_format: TextFormat,
        **kwargs
        ):

        if kwargs.get('outdir'):
            out_dir = kwargs.get('outdir')
        else:
            out_dir = txt_transliterand.parent

        def transliterate():
            ...
        # out_filename = f"{txt_transliterand.stem}-{source_format}-to-{target_format}-transliterated.txt"
        # out_path = out_dir.joinpath(out_filename)
        # with (
        #     open(txt_transliterand, 'r+') as opened_source_file,
        #     open(out_path, 'w+') as opened_target_file
        # ):

        #     for line in opened_source_file:
        #         found_hulq_words = source_kp.extract_keywords(line)
        #         found_eng_words = eng_kp.extract_keywords(line)

        #     if len(found_hulq_words) > len(found_eng_words) or len(found_eng_words) == 0:
        #         transl_line = repl.transliterate_string(line, 
        #         source_format, 
        #         target_format)

        #         opened_target_file.write(line)


class Transliterator():

    def __init__(
            self,
            text_counter: TextCounter,
            graphemes: Graphemes,
            wordlists=None
    ):
        self.text_counter = text_counter
        self.graphemes = graphemes
        self.docx_transliterator = DocxTransliterator(self.text_counter, self.graphemes)
        self.txt_transliterator = TxtTransliterator(self.text_counter, self.graphemes)

    def transliterate_string(
        self,
        source_string: str,
        source_format: TextFormat,
        target_format: TextFormat,
        ):
        """just transliterates a single string.
        It's a thing of beauty"""

        transliterated_string = repl.transliterate_string(
            source_string,
            source_format,
            target_format,
            self.graphemes)
        return transliterated_string

    def transliterate_file(self, _file: TransliterandFile):
        if _file.file_type == ".docx":
            return self.docx_transliterator.transliterate(_file)
        elif _file.file_type == ".txt":
            return self.txt_transliterator.transliterate_txt_wordlist(_file)

    # def update_new_wordlist(
    #     source_format: TextFormat,
    #     current_keywordprocessor: KeywordProcessor,
    #     new_keywords: list,
    #     keywordprocessors_dict: dict):
    #     """update the running wordlist with things that were transliterated

    #     Arguments:
    #         source_format -- the current source format (to get base wordlist folder)
    #         current_keywordprocessor -- the current keyword processor (to get the keywords)
    #         new_keywords -- all the found 
    #     """
    #     working_wordlist_filepath = Path(__file__).parent / ('resources/wordlists/new-hulq-wordlist-' +
    #                                                     source_format.to_string() +
    #                                                         '.txt')

    #     current_keywords = current_keywordprocessor.get_all_keywords()
        
    #     set(new_keywords.get_all_keywords().keys())


class TransliterandFileHandler(FileHandler):
    """Class to hold all files to transliterate and transliterates them upon
    request.
    """
    def __init__(
            self,
            files_list: list[Path | str | FileStorage],
            transliterator: Transliterator,
            out=Path | None,
            **kwargs):

        super().__init__(files_list)

        self.search_method = kwargs.get('search_method')
        self.transliterand_files = [
            TransliterandFile(
                _file,
                source_format=kwargs.get("source_format"),
                target_format=kwargs.get("target_format"),
                search_method=self.search_method,
                out=out
                )
            for _file in self.files
            ]
        self.transliterator = transliterator

        if (source_format := kwargs.get("source_format")):
            for file in self.transliterand_files:
                file.source_format = source_format
        if (target_format := kwargs.get("target_format")):
            for file in self.transliterand_files:
                file.target_format = target_format

    def transliterated(self) -> list[Path]:

        return [
            self.transliterator.transliterate_file(_file)
            for _file in self.transliterand_files
            ]


if __name__ == "__main__":
    ...