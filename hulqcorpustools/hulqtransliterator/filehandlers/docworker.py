'''
takes either a .doc or .docx file and transliterates based on font
or style
TODO: write real tests and write test checking
    - for each file format to another
'''


from docx import Document as load_docx
from docx.document import Document
from docx.text import paragraph, run
from docx import table
from flashtext import KeywordProcessor
from pathlib import Path
import os

from hulqcorpustools.utils.keywordprocessors import HulqKeywordProcessors
from ...resources.constants import FileFormat, GraphemesDict, TransliterandFile
from ..transliterator import replaceengine as repl

class DocxTransliterator():



    def transliterate_docx_wordlist(
        transliterand: Path,
        source_format = FileFormat,
        target_format = FileFormat,
        kps = HulqKeywordProcessors,
        **kwargs):
        '''transliterates an entire docx file by wordlist

        Keyword arguments:
            update_wordlist -- include this if you want to update the wordlist
                            with words found in the transliterated lines
            font_search -- include this 
        '''

        document = load_docx(transliterand)
        out_filename = f'{transliterand.stem} {source_format.to_string()} to {target_format.to_string()} transliterated.docx'

        if (out_dir := kwargs.get('outdir')): # type: Path
            out_path = out_dir.joinpath(out_filename)
        else:
            out_path = transliterand.parent.joinpath(out_filename)
        for par in document.paragraphs:
            par_text_parts = par.text.split('\t')
            new_par_text_parts = []
            for par_text_part in par_text_parts:
                if kps.determine_language_from_text(par_text_part) != 'english':
                    par_text_part = repl.transliterate_string_replace(
                        par_text_part,
                        source_format,
                        target_format
                    )
                new_par_text_parts.append(par_text_part)
            # for par_text_part in par_text_parts:
            #     if kps.determine_language_from_text(par_text_part) != 'english':
            #         par_text_part = repl.transliterate_string_replace(
            #             par_text_part,
            #             source_format,
            #             target_format
            #         )

            par.text = '\t'.join(new_par_text_parts)
            ...

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if kps.determine_language_from_text(cell.text) != 'english':
                        cell_text = repl.transliterate_string_replace(
                            cell.text,
                            source_format,
                            target_format
                        )
                        cell.text = cell_text

        document.save(out_path)
        return out_path 

    def transliterate_docx_font(
        transliterand: Path,
        source_format = FileFormat,
        target_format = FileFormat
        ):

        document = load_docx(transliterand)
        out_filename = f'{transliterand.stem} {source_format.to_string()} to {target_format.to_string()} transliterated.docx'
        out_path = transliterand.parent.joinpath(out_filename)

        for par in document.paragraphs:
            if par.style.font.name == 'Straight':
                par.text = repl.transliterate_string_replace(
                    par.text,
                    source_format,
                    target_format)
            else:
                for run in par.runs:
                    if run.font.name == 'Straight':
                        run.text = repl.transliterate_string_replace(
                            run.text,
                            source_format,
                            target_format
                        )

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.name == 'Straight':
                                run.text = repl.transliterate_string_replace(
                                    run.text,
                                    source_format,
                                    target_format
                                )


        document.save(out_path)
        return(out_path)

    def update_new_wordlist(
        source_format: FileFormat,
        current_keywordprocessor: KeywordProcessor,
        new_keywords: list,
        keywordprocessors_dict: dict):
        """update the running wordlist with things that were transliterated

        Arguments:
            source_format -- the current source format (to get base wordlist folder)
            current_keywordprocessor -- the current keyword processor (to get the keywords)
            new_keywords -- all the found 
        """
        working_wordlist_filepath = Path(__file__).parent / ('resources/wordlists/new-hulq-wordlist-' +
                                                        source_format.to_string() +
                                                            '.txt')

        current_keywords = current_keywordprocessor.get_all_keywords()
        
        set(new_keywords.get_all_keywords().keys())


    
if __name__ == "__main__":
    ...