
import itertools
import os
from pathlib import Path
import re
from typing import Optional, Generator
from zipfile import BadZipFile

import docx

class CorpusDocx(docx.document.Document):
    """class for each corpus docx file

    Arguments:
        docx -- _description_
    """
    def __init__(self, docx_path: Path):
        """init docx file 

        Arguments:
            docx_path -- path to a docx file
        """
        self.filename = docx_path.name
        self.stem = docx_path.stem
        self.versionless_stem = self._get_version(self.filename)[0]
        self.version = self._get_version(self.filename)[1]
        
        try:
            self.Document = docx.Document(docx_path)
        except BadZipFile as error:
            print(error)
            print(self.filename, 'is not a good docx file despite suffix')

    @classmethod
    def _get_version(
        cls,
        _docx_filename
        ):
        _version_re = re.compile('(.*)([A-Z]{2,4})([1-9]{1,4})(?:\.docx|$)')
        _version_parse = _version_re.search(_docx_filename)
        _versionless_stem = _version_parse[1]
        _version = F'{_version_parse[2]}{_version_parse[3]}'
        return (_versionless_stem, _version)

    def _get_paragraph_text(
        self
        ):
            return (paragraph.text + '\n' for paragraph in self.Document.paragraphs)

    def _get_paragraph_style(
        self
        ):
            return (paragraph.style.name for paragraph in self.Document.paragraphs)

    def _get_paragraph_comments(
        self
        ):

        return [[comment.text for comment in paragraph.comments] for paragraph in self.Document.paragraphs]
        

if __name__ == "__main__":
    corpus_docx_list = list(CorpusDocx(_corpus_docx) for _corpus_docx in Path(os.environ.get('CORPUS_TEST_DATA_FOLDER')).glob('*.docx'))

    print(corpus_docx_list[0].Document.__getattribute__)

    print(CorpusDocx().__)
    
    # print([_docx.filename for _docx in corpus_docx_list])
    # first_docx = corpus_docx_list.__next__()
    # print(first_docx.filename)
    # text_style = zip(
    # first_docx._get_paragraph_text(),
    # first_docx._get_paragraph_style())
    # for i in text_style:
    #     print(i)
        
    pass