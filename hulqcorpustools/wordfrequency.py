
from collections import Counter
from collections.abc import Iterable
import json
from pathlib import Path
import re

from docx import Document as init_docx
from docx.document import Document

from hulqcorpustools.utils.files import FileHandler
from hulqcorpustools.utils.keywordprocessors import kp
from hulqcorpustools.resources.constants import TextFormat

from werkzeug.datastructures import FileStorage

reg_str_pattern = re.compile(r"[\"\'\!\.\,\?\[\]\(\)\“\”\;]|LH\t|LE\t")


class WordCounter():

    def __init__(self):
        self.total = Counter()

    def iter_count_words(self, line_list: Iterable[str]):
        for line in line_list:
            if kp.determine_text_format(line) != TextFormat.ENGLISH:
                words = (reg_str_pattern.sub("", word) for word in line.split())
                self.total.update(words)

    def count_words(self, _str: str):
        str_lines = _str.split("\n")
        self.iter_count_words(str_lines)


class FileWordCounter(WordCounter):

    def __init__(self, files_list: list[FileStorage]):
        super().__init__()
        self.file_list = FileHandler(files_list)

    def count_txt_words(self):
        for _file in self.file_list.txt_files:
            with open(_file) as _txt:
                self.iter_count_words(_txt)

    def count_docx_words(self):
        par_text_lines = []
        for _file in self.file_list.docx_files:
            _docx = init_docx(_file)  # type: Document
            for _par in _docx.paragraphs:
                par_text_lines.append(_par.text)

            for table in _docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for _par in cell.paragraphs:
                            par_text_lines.append(_par.text)

        self.iter_count_words(par_text_lines)

    def count_file_words(self):
        self.count_txt_words()
        self.count_docx_words()





def write_frequency_to_txt(output_filepath: Path, counted_words: Counter):
    """writes the result of counting word frequency to a txt file

    Arguments:
        output_filepath -- path to the output file
        counted_words -- output of the counted_words fn
    """
    with open(output_filepath, 'w+') as output_file:
        for i in counted_words.most_common():
            output_file.write(i[0] + '\t' + str(i[1]) + '\n')

def write_frequency_to_json(output_filepath: Path, counted_words: Counter):
    """writes the reuslt of counting word frequency to a json

    Arguments:
        output_filepath -- path to the output file
        counted_words -- output of the counted_words fn
    """
    with open(output_filepath, 'w+') as open_file:
        json.dump(dict(counted_words.most_common()), open_file, ensure_ascii=False, indent=0)
