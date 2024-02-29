
from importlib import resources
from inspect import Attribute
from pathlib import Path
import re

from aenum import Enum, Flag, auto
import docx
from docx.enum.style import WD_STYLE_TYPE
from zipfile import BadZipFile

from . import docxannotationdata
from .constants import TextFormat
from ..utils.languagesusser import determine_language_from_text

class _HulqAnnotationStyles():

    def __init__(
        self
        ):

        self._annotation_template_docx = docx.Document(
                                resources.open_binary(
                        docxannotationdata,"hulq-story-annotation-template.docx"
                        ))
        self._docx_annotation_styles = self._annotation_template_docx.styles # type: docx.styles.styles.Styles

        self.annotation_style_names = {
            "Title",
            "Story author name",
            "Notes - Biographical notes",
            "Notes - Collection bibliographic notes",
            "Notes - Story bibliographic notes",
            "Notes - Advice needed"
            "Notes - Story inline notes",
            "Notes",
            "Hul’q’umi’num’ line numbering",
            "Text line - TBD",
            "Text line - Hul’q’umi’num’ line",
            "Text line - English translation of Hul’q’umi’num’",
            "Text line - English spoken only",
            "Text line - Hul’q’umi’num’ line with no translation",
            "Text line - English with no Hul’q’umi’num’",
            "Text - Mention of Hul’q’umi’num’ word in Hul’q’umi’num’",
            "Text - English word in Hul’q’umi’num’ line",
            "Text - Hul’q’umi’num’ word in English translation",
            "Text - Speaker switch",
            "Heading 1",
            "Heading 2",
            "Heading 3"
        }

        self._heading_styles = [
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Title",
            "Story author name"
        ]

        self._hulq_line_styles = [
            "Text line - Hul’q’umi’num’ line",
            "Text line - Hul’q’umi’num’ line with no translation"
        ]

        self._eng_line_styles = [
            "Text line - English translation of Hul’q’umi’num’",
            "Text line - English spoken only",
            "Text line - English with no Hul’q’umi’num’"
        ]

        self._grouped_styles = self._hulq_line_styles + self._eng_line_styles + ["Text line - TBD"]

        self._character_styles = [
            "Text - Mention of Hul’q’umi’num’ word in Hul’q’umi’num’",
            "Text - English word in Hul’q’umi’num’ line",
            "Text - Hul’q’umi’num’ word in English translation",
            "Text - Speaker switch",
        ]

        self.annotation_styles = {
        _HulqStyle.name : _HulqStyle for _HulqStyle in self._docx_annotation_styles if _HulqStyle.name in self.annotation_style_names
        }

    def _validate(
        self,
        _annotation_style: docx.styles.style._ParagraphStyle |
                            docx.styles.style._CharacterStyle
    ):

    
        try:
            _annotation_valid = self.annotation_styles.get(_annotation_style.name)
            if _annotation_valid and _annotation_valid.name in self.annotation_style_names:
                return _annotation_valid
            else:
                return None
        except TypeError as e:
            print(f'oh no...{e} error with {_annotation_valid}')
    

HulqAnnotationStyles = _HulqAnnotationStyles()

class _HulqAnnotationStyle():

    def __init__(
        self,
        _annotation_style: docx.styles.style._ParagraphStyle,
        verbose=False
        ):

        # todo work on this
        self._annotation_style = self._validate(_annotation_style)
        if not self._annotation_style:
            _annotation_style = None


    @classmethod
    def _validate(cls, _potential_annotation_style):
        if HulqAnnotationStyles._validate(_potential_annotation_style):
            return _potential_annotation_style
        else:
            
            return None

class _DocxCorpusPar():

    # TODO: deal with multiple runs in one line, e.g. character styles
    # if len(_text_par.runs) > 1:
    #     for i in _text_par.runs:
    #         print(i.style.name)

    def __init__(
        self,
        _text_par
        ):

        self._text_par = _text_par
        self._text = _text_par.text.strip()

        if len(self._text) < 1:
            _text_par.delete()

        # self._language = None # needed for grouping lines
        # self._annotation_style = None

        def _process_line_text(self):
            _clear_run_formatting(self)
            _valid_annotation_style = HulqAnnotationStyles._validate(_text_par.style)

            if _valid_annotation_style is not None:
                # if valid: style has been manually applied and so assumed OK
                self._annotation_style = _valid_annotation_style
                self._language = None

                return

            else:
                _set_interim_annotation_style(self)

        def _set_interim_annotation_style(self):
            """in the case that the lines of hulq or english have not been 
            annotated (as they are most numerous and the most time-consuming to
            annotate)
            
            """
            # if not valid: it might be a line of text in hulq or english;
            # determine this
            _language = determine_language_from_text(_text_par.text)
            
            if _language is None:
                # if language can't be determined: needs to be fixed by hand
                _annotation_style = HulqAnnotationStyles.annotation_styles.get('Notes - Advice needed')
            else:
                # if language is determined: the next step is to hand it to
                # the text line group to see whether it is
                _annotation_style = HulqAnnotationStyles.annotation_styles.get('Text line - TBD')

            self._language = _language
            self._annotation_style = _annotation_style

        def _clear_run_formatting(self):
            """if there is any kind of extraneous formatting directly applied:
            clear it
            """
            for run in self._text_par.all_runs:
                run.font.bold = False
            # for run in self._text_par.runs:
            #     print(run)

        _process_line_text(self)

class _DocxCorpusParGroup():
    """a group of text lines associated with one another-- specifically, a line
    of hulq and a line of english, but in some cases there is a singlet
    """

    def __init__(
        self,
        hulq_line = None | _DocxCorpusPar,
        eng_line = None | _DocxCorpusPar,
        # _notes = None | list[_DocxCorpusPar]
        ):
        ...
        self.hulq_line = hulq_line
        self.eng_line = eng_line
        # self._notes = _notes

        self.current_line_languages = {
            TextFormat.HULQ_FORMATS : None,
            TextFormat.ENGLISH : None
        }


    def _update_group(self,
            _incoming_line: _DocxCorpusPar,
            ):
        """groups line of hulq and line of english together and applies style,
        or appropriately applies style 

        Arguments:
            _incoming_line -- the new line to update the group with
        """

        _incoming_style = _incoming_line._text_par.style
        _incoming_language = _incoming_line._language

        # if style of line coming has been annotated already, but isn't hulq or english line: skip
        if _incoming_style.name in HulqAnnotationStyles.annotation_style_names and \
        _incoming_style.name not in HulqAnnotationStyles._grouped_styles:
            self._finish_group()
            return

        # something weird has happened - set that it needs to be fixed manually
        if _incoming_language == None:
            _incoming_line.style = HulqAnnotationStyles.annotation_styles.get('Notes - Advice needed')
            return

        # normalize incoming language if hulq from specific (e.g. orthography)
        # to general
        if type(_incoming_language) == TextFormat and \
        _incoming_language & TextFormat.HULQ_FORMATS:
            _incoming_language = TextFormat.HULQ_FORMATS


        # if there is any hulq: it is always a new line
        if _incoming_language == TextFormat.HULQ_FORMATS and \
            self.current_line_languages.get(_incoming_language):
            self._finish_group()

        # if the last line has not had one of hulq or english already: update
        # that one is accounted for and continue
        if not self.current_line_languages.get(_incoming_language):
            self.current_line_languages.update(
                {_incoming_language: _incoming_line}
                )

        # if the last line has one of the same language: set their styles
        else:
            self._finish_group()

    def _finish_group(self):
        """returns everything in the line now that no new info is added

        Returns:
            finished line
        """
        _finish_hulq = self.current_line_languages.get(TextFormat.HULQ_FORMATS)
        _finish_eng = self.current_line_languages.get(TextFormat.ENGLISH)

        # if _finish_hulq:
        #     _finish_hulq._text_par.
        if _finish_hulq and not _finish_eng:
            _finish_hulq._text_par.style = \
                HulqAnnotationStyles.annotation_styles.get("Text line - Hul’q’umi’num’ line with no translation")

        elif not _finish_hulq and _finish_eng:
            _finish_eng._text_par.style = \
                HulqAnnotationStyles.annotation_styles.get("Text line - English with no Hul’q’umi’num’")
                

        if _finish_hulq and _finish_eng:
            _finish_hulq._text_par.style = \
                HulqAnnotationStyles.annotation_styles.get("Text line - Hul’q’umi’num’ line")
            _finish_eng._text_par.style = \
                HulqAnnotationStyles.annotation_styles.get("Text line - English translation of Hul’q’umi’num’")


        self.current_line_languages.update(
            {TextFormat.HULQ_FORMATS: None,
            TextFormat.ENGLISH: None}
        )

class CorpusDocx(docx.document.Document):
    """class for each corpus docx file

    Arguments:
        docx -- _description_
    """
    def __init__(self,
        docx_path: Path,
        **kwargs):
        """init docx file 

        Arguments:
            docx_path -- path to a docx file
        """
        self.filename = docx_path.name

        try:
            self._Document = docx.Document(docx_path)
        except BadZipFile as error:
            print(error)
            print(self.filename, 'is not a good docx file despite suffix')

        self.stem = docx_path.stem
        self.versionless_stem = self._get_version(self.filename)[0]
        self.version = self._get_version(self.filename)[1]
        

        if kwargs.get('corpusformat') is True:

            self.update_formatting_styles()
            self.formatted_path = docx_path.parent.parent / 'output'
            self.formatted_filename = self.formatted_path / f"{self.filename} corpus formatted.docx"
            self.process_paragraphs(self._Document.paragraphs)
            self._Document.save(self.formatted_filename)

    def update_formatting_styles(self):
        """updates the doc with styles for corpus formatting
        """

        docx_style_names = [style.name for style in self._Document.styles]


        for _anno_style in HulqAnnotationStyles.annotation_styles.values():

            if _anno_style.name not in docx_style_names:
                
                _new_anno_style = self._Document.styles.add_style(_anno_style.name, _anno_style.type)

            # self._Document.styles[_anno_style.name].element


            # print('***_element v', _new_anno_style._element.xml, '***_element ^', sep='\n')
            # print('@@@ element v', _new_anno_style.element.xml, '@@@ element ^', sep='\n')

            self._Document.styles[_anno_style.name].element = _anno_style.element

            print(_anno_style.element.xml)
            # print(_anno_style.element.rPr)

        # for docx_style in self._Document.styles:
        #     print(docx_style.element.xml)

    def delete_blank_paragraph(
        self,
        _paragraph,
        ):
        if len(_paragraph.text) < 1:
            _paragraph.delete()

    def pars_into_groups(
        self,
        _paragraph: _DocxCorpusPar
        ):

        self.par_group._update_group(_DocxCorpusPar(_paragraph))

    def process_paragraphs(
        self,
        _paragraphs):
        
        self.par_group = _DocxCorpusParGroup()
        for paragraph in _paragraphs:
            if len(paragraph.text) < 1:
                paragraph.delete()
                continue
            self.pars_into_groups(paragraph)

    def _get_version(
        cls,
        _docx_filename
        ):
        _version_re = re.compile('(.*)([A-Z]{2,4})([1-9]{1,4})(?:\.docx|$)')
        _version_parse = _version_re.search(_docx_filename)

        if not _version_parse:
            return(None, None)
        _versionless_stem = _version_parse[1]
        _version = F'{_version_parse[2]}{_version_parse[3]}'

        return (_versionless_stem, _version)

    def _get_paragraph_text(
        self
        ):
            return (paragraph.text + '\n' for paragraph in self.Document.paragraphs)

    def _get_paragraph_style(
        self
        ):
            return (paragraph.style.name for paragraph in self.Document.paragraphs)

    def _get_paragraph_comments(
        self
        ):

        return [[comment.text for comment in paragraph.comments] for paragraph in self.Document.paragraphs]

