
import csv
import json
from importlib import resources as res
import os
from pathlib import Path
import re

from docx import Document as construct_doc
from docx.document import Document
from docx.text import paragraph, run

from hulqcorpustools.utils.keywordprocessors import HulqKeywordProcessors
from hulqcorpustools.resources.constants import TextFormat

kp = HulqKeywordProcessors(eng=True)


class _CorpusMetadata():

    authors = res.files('metadata').joinpath('.authors')
    authors_list = open(authors).readlines()
    authors_codes = {
        i.split('\t')[1].strip(): 
        i.split('\t')[0]
        for i in authors_list}

    story_codes = json.load(open(res.files('metadata').joinpath('.storycodes.json'))) # type: dict[dict]

    def _generate_storycode_json():
        story_codes = res.files('metadata').joinpath('.storycodes.csv')

        with open(story_codes, 'r', newline='', encoding="utf-8-sig") as story_csv:
            stories_dict = dict()
            story_reader = csv.DictReader(story_csv)
            
            with open(res.files('metadata').joinpath('.storycodes.json'), 'w', encoding="utf-8") as story_json_file:
                for row in story_reader:
                    story_entry = {row.get('story_id'): row}
                    stories_dict.update(story_entry)
                
                json.dump(stories_dict, story_json_file, ensure_ascii=False, indent=2)
                story_json_file.write('\n')



class CorpusStory():

    def __init__(self,
                 metadata: dict):
        for i, v in metadata.items():
            setattr(self, i, v)
        self.pars = []

class CorpusAuthorCollection():

    def _get_stories(self, _docx: Document):
        _current_story_pars = []
        _current_line_group = {
        'hulq': None,
        'eng': [],
        'number': 1}
        _current_story = None

        def _add_line_number(_line_group: dict):
            text_abbreviation = f"{_current_story.story_id}.{_current_line_number}"
            if len(_line_group.get('eng')) > 1:
                _final_line = _line_group.get('eng')[-1]
            elif len(_line_group.get('eng')) == 1:
                _final_line = _line_group.get('eng')[0]
            else:
                _final_line = _line_group.get('hulq')
            _final_line.add_run(f"\n{text_abbreviation}")

            _line_group.update({'hulq':_par})
            _line_group.update({'eng': []})
            _line_group.update({'number': _current_line_number + 1})

        for _par in _docx.paragraphs:
            if len(_par.text) == 0:
                continue

            if _par.style.name == 'Heading 1':
                continue

            if _par.style.name == "Heading 2":
                if _current_line_group.get('hulq'):
                    _current_line_number = _current_line_group.get('number')
                    _add_line_number(_current_line_group)
                _story_title = self._get_story_title(_par)
                _story_metadata = self._get_story_metadata(_story_title)
                try:
                    _current_story = CorpusStory(_story_metadata)
                except AttributeError:
                    print(_story_title)
                _current_line_group.update({
                    'hulq': None,
                    'eng': [],
                    'number': 1
                })

            elif _par.text[0:2] in ['N\t', 'A\t', 'C\t']:
                continue

            elif _par.text[0:2] == 'E\t':
                _par.text = _par.text.removeprefix('E\t')
                if _current_line_group.get('hulq'):
                    _current_eng_lines = _current_line_group.get('eng') # type: list
                    _current_eng_lines.append(_par)
                    _current_line_group.update({'eng':_current_eng_lines})


            elif _par.text[0:2] == 'L\t':
                _par.text = _par.text.removeprefix('L\t')
                _current_line_number = _current_line_group.get('number')
                if _current_line_group.get('hulq') is None:
                    _current_line_group.update({'hulq' : _par})
                else:
                    _add_line_number(_current_line_group)



            elif _par.style.name != "Heading 2" and self._check_language(_par) == 'hulq':
                _current_line_number = _current_line_group.get('number')

                if _current_line_group.get('hulq'):
                    _add_line_number(_current_line_group)

                else:
                    _current_line_group.update({'hulq':_par})



            elif self._check_language(_par) == 'eng' or _par.text[0:2] == 'E\t':
                if _current_line_group.get('hulq'):
                    _current_eng_lines = _current_line_group.get('eng') # type: list
                    _current_eng_lines.append(_par)
                    _current_line_group.update({'eng':_current_eng_lines})

    def _get_story_title(self, _story_title_par: paragraph.Paragraph):
        _story_title = _story_title_par.text.strip()
        _story_title.replace(' \t', '\t') #  _par.text[0] == "T":
        _story_title = _story_title.removeprefix('T\t')
        _story_title = _story_title.removeprefix('Title\t')
        _story_title = _story_title.strip()
        return _story_title

    def _get_story_metadata(self, _title: str):

        def _check_for_title_match(_sub_str: str):
            if len(_sub_str) > 1 and _sub_str.casefold() in _title.casefold():
                return True

        for story_index in _CorpusMetadata.story_codes.items():
            story_index_author_code = story_index[0].split('.')[0]
            story_index_metadata = story_index[1]
            if story_index_author_code == self.author_code:
                eng_lookup_title = story_index_metadata.get('story_title_eng')
                hulq_lookup_title = story_index_metadata.get('story_title_hulq')
                if _check_for_title_match(eng_lookup_title) or _check_for_title_match(hulq_lookup_title):
                    return story_index[1]
                else:
                    continue

    def _compile_single_story(self, _current_story, _pars: list[paragraph.Paragraph]):
        print(_current_story.title)
        for i in _pars:
            ...

    def _check_language(self, _par: paragraph.Paragraph):
        _current_line_lang = kp.determine_language(_par.text)

        if _current_line_lang == TextFormat.APAUNICODE or _current_line_lang == TextFormat.ORTHOGRAPHY:
            return 'hulq' 
        elif _current_line_lang == 'english':
            return 'eng'

    def _get_author_name(self, _text_file_name: Path):
        repl_underscore_trans = str.maketrans({"_": " "})
        author_name = _text_file_name.stem
        author_name = author_name.removeprefix('Text_corpus_')
        author_name = author_name.translate(repl_underscore_trans)
        author_name = author_name.split('-')[0]
        return author_name

    trim_title_prefix = lambda x: re.sub('T.*\t', '', x)

    def __init__(self, _docx_file: Path):
        self.docx_file_path = _docx_file
        self.docx_numbered_file_name = f"{self.docx_file_path.stem} numbered.docx"
        self.docx = construct_doc(_docx_file) # type: Document
        self.author_name = self._get_author_name(self.docx_file_path)
        self.author_code = _CorpusMetadata.authors_codes.get(self.author_name)
        self.collection_stories = []
        print(self.docx_numbered_file_name)

        self._get_stories(self.docx)

def number_all_corpus_docx(corpus_docx_dir: Path):
    docx_input_dir = corpus_docx_dir / 'input'
    docx_saved_dir = corpus_docx_dir / 'numbered'
    for i in docx_input_dir.iterdir():
        if i.name[0] in ["~", "."]:
            continue
        collection_docx = CorpusAuthorCollection(i)
        collection_docx.docx.save(docx_saved_dir / collection_docx.docx_numbered_file_name)


def make_compiled_corpus_docx(corpus_docx_dir: Path):
    compiled_docx = construct_doc() # type: Document
    for docx_file in Path(corpus_docx_dir / "input").iterdir():
        if docx_file.name[0] in ['.', '~']:
            continue
        constituent_doc = construct_doc(docx_file) #type: Document

        for _par in constituent_doc.paragraphs:
            compiled_docx._body.
            compiled_docx.paragraphs.append(_par)

    final_compiled_docx_path = Path(corpus_docx_dir / 'compiled/compiled_corpus.docx')
    compiled_docx.save(final_compiled_docx_path)

if __name__ == "__main__":

    corpus_docx = Path(os.environ['CORPUSOUT'])

    make_compiled_corpus_docx(corpus_docx)